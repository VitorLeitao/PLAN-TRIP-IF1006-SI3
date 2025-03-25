from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.shared import RGBColor
import importlib
import re
from datetime import datetime
from bs4 import BeautifulSoup

class WordTemplate:
    
    def __init__(self, margem_esquerda=0.5, margem_direita=0.5):
        self.doc = Document()
        self.set_margens(margem_esquerda, margem_direita)

    def set_margens(self, margem_esquerda, margem_direita):
        section = self.doc.sections[0]
        section.left_margin = Inches(margem_esquerda)
        section.right_margin = Inches(margem_direita)
    
    def add_footer(self):
        section = self.doc.sections[0]
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]

        data = datetime.now().strftime("%d/%m/%Y às %H:%M")
        footer_paragraph.text = f"Gerado por IA em {data}"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Adiciona o número da página atual
        self.add_field_footer(footer_paragraph, "PAGE")
        footer_paragraph.add_run(" de ")
        # Adiciona o número de páginas
        self.add_field_footer(footer_paragraph, "NUMPAGES")


    def add_field_footer(self, paragraph, field_type):
        run_page_num = paragraph.add_run()
        fld_char1 = OxmlElement('w:fldChar')
        fld_char1.set(ns.qn('w:fldCharType'), 'begin')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(ns.qn('xml:space'), 'preserve')
        instr_text.text = field_type
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(ns.qn('w:fldCharType'), 'end')

        run_page_num._r.append(fld_char1)
        run_page_num._r.append(instr_text)
        run_page_num._r.append(fld_char2)

    def add_linha(self, width_pct=100):
        p = self.doc.add_paragraph()
        p_pr = p._element.get_or_add_pPr()

        top_border = OxmlElement('w:top')
        top_border.set(qn('w:val'), 'single') 
        top_border.set(qn('w:sz'), '4')
        p_pr.append(top_border)

        if width_pct < 100:
            p_format = p.paragraph_format
            p_format.space_before = Pt(100)
            total_width = 100
            margin_left = (total_width - width_pct) / 2
            p_format.left_indent = Inches(margin_left / 100 * 6.5)
            p_format.right_indent = Inches(margin_left / 100 * 6.5)

    def handle_p(self, element, align):
        p = self.doc.add_paragraph()
        p.alignment = align
        for sub_element in element.children:
            run = self.handle_run(p, sub_element)

    def handle_h(self, element, align):
        level = 1 if element.name == "h1" else 2
        fonte = 22 if element.name == "h1" else 18
        heading = self.doc.add_heading(element.get_text(), level=level)
        heading.alignment = align
        run = heading.runs[0]
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(fonte)

    def handle_hr(self, element):
        width_pct = 100
        if 'style' in element.attrs:
            style = element.attrs['style']
            match = re.search(r'width:\s*(\d+)%', style)
            if match:
                width_pct = int(match.group(1))
        self.add_linha(width_pct)

    def handle_ul(self, element):
        for li in element.find_all('li'):
            p = self.doc.add_paragraph(style='List Bullet')
            for sub_element in li.children:
                self.handle_run(p, sub_element)

    def handle_run(self, paragraph, sub_element):
        run = paragraph.add_run(sub_element.get_text())
        if sub_element.name == "strong":
            run.bold = True
        elif sub_element.name == "i":
            run.italic = True
        run.font.name = 'Arial'

    def processar_elemento(self, element, align=None):
        if element.name == "p":
            self.handle_p(element, align)
        elif element.name in ["h1", "h2"]:
            self.handle_h(element, align)
        elif element.name == "ul":
            self.handle_ul(element)
        elif element.name == "hr":
            self.handle_hr(element)

    async def convert(self, q, content) -> Document:
        content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
        content = content.replace("*", "-")
        
        self.doc = importlib.import_module(f'templates.{q.client.selected_relatorio}.{q.client.selected_relatorio}_service').get_header_doc(self.doc, q)
        
        soup = BeautifulSoup(content, "html.parser")
        soup_content = list(soup.children)[0] # Pegando apenas o conteudo dentro da DIV pai
        for element in soup_content:
            if element.name == "div":
                for sub_element in element.children:
                    self.processar_elemento(sub_element, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                self.processar_elemento(element, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        self.add_footer()

        return self.doc
